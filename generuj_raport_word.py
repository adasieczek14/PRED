# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r'C:\Users\admin\Desktop\PRACA INZYNIERSKA\Analiza_Strategii_Typowania.docx'

# Kolory jako krotki (R, G, B) - NIE jako RGBColor (ma inne atrybuty wewnatrz)
C_DARK   = (0x1E, 0x1E, 0x2E)
C_GREY   = (0x60, 0x60, 0x60)
C_WHITE  = (0xFF, 0xFF, 0xFF)
C_LGREY  = (0xF2, 0xF2, 0xF2)
C_LGREEN = (0xE6, 0xF4, 0xEA)
C_LRED   = (0xFD, 0xED, 0xED)
C_LYELL  = (0xFF, 0xF9, 0xE6)
C_RED    = (0xC0, 0x20, 0x20)
C_ORANGE = (0xD4, 0x6F, 0x00)
C_NAVYBL = (0x1A, 0x3A, 0x5C)
C_DKGREEN= (0x10, 0x6A, 0x35)

def rgb(t): 
    return RGBColor(t[0], t[1], t[2])

def hex_color(t):
    return '{:02X}{:02X}{:02X}'.format(t[0], t[1], t[2])

def set_cell_bg(cell, color_tuple):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(color_tuple))
    tcPr.append(shd)

def set_borders(cell, color='AAAAAA'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(doc, text='', bold=False, italic=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = rgb(color)
    return p

def add_head(doc, text, level=2, color=C_DARK):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = rgb(color)
    r.font.size = Pt(16 if level == 1 else (13 if level == 2 else 11))

def make_table(doc, headers, rows, col_widths=None, hdr_bg=C_DARK, row_colors=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Naglowek
    hdr_cells = table.rows[0].cells
    for i, (h, cell) in enumerate(zip(headers, hdr_cells)):
        set_cell_bg(cell, hdr_bg)
        set_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = rgb(C_WHITE)
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Dane
    for ri, row_data in enumerate(rows):
        bg = row_colors[ri] if (row_colors and ri < len(row_colors)) else (C_LGREY if ri % 2 == 0 else C_WHITE)
        row = table.add_row()
        for ci, (val, cell) in enumerate(zip(row_data, row.cells)):
            set_cell_bg(cell, bg)
            set_borders(cell, 'CCCCCC')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.bold = (ci > 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

# ============================================================
# BUDOWANIE DOKUMENTU
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# TYTUL
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(4)
r = t.add_run('ANALIZA STRATEGII TYPOWANIA')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = rgb(C_DARK)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.paragraph_format.space_after = Pt(2)
r2 = t2.add_run('System Predykcji KNN & XGBoost — Co grac, czego unikac')
r2.font.size = Pt(12); r2.italic = True; r2.font.color.rgb = rgb(C_GREY)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.paragraph_format.space_after = Pt(14)
r3 = t3.add_run('Raport oparty na danych realnych (nie backtest): 11.03.2026 - 10.04.2026')
r3.font.size = Pt(10); r3.font.color.rgb = rgb(C_GREY)

# ---- 1. KONTEKST PROBY ----
add_head(doc, '1. Podstawa analizy — wielkosc proby', level=2)
add_para(doc,
    'Analiza opiera sie wylacznie na meczach zebranych i obstawianych na zywo — '
    'nie jest to backtest. Wazne: KNN zbiera mecze od 11 marca (31 dni), '
    'natomiast XGBoost startuje od 25 marca (tylko 16 dni) — dlatego probe XGB nalezy traktowac jako mniejsza.',
    italic=True, color=C_GREY, size=10, sa=8)

make_table(doc,
    ['Model', 'Okres dzialania', 'Liczba meczow OGOLNIE', 'Win Rate ogolny', 'ROI ogolny (Flat 100 zl)'],
    [
        ['KNN',     '11.03.2026 - 10.04.2026  (31 dni)', '1 237', '52.8%', '-24 917 zl'],
        ['XGBoost', '25.03.2026 - 10.04.2026  (16 dni)', '645',   '57.2%', '-6 973 zl'],
    ],
    col_widths=[2.2, 5.0, 4.0, 3.2, 4.2],
    row_colors=[C_LGREEN, C_LYELL]
)
add_para(doc,
    '\nOgolny ROI jest ujemny, bo system typuje wszystkie mecze bez filtra progu pewnosci. '
    'Dopiero po filtrowaniu wyniki diametralnie sie zmieniaja — patrz ponizej.',
    italic=True, color=C_ORANGE, size=9.5, sa=14)

# ---- 2. SKUTECZNOSC PER PROG ----
add_head(doc, '2. Skutecznosc wg progu pewnosci modelu', level=2)
add_para(doc, 'Ponizsza tabela pokazuje, ze typowanie z filtrem progu jest absolutnie kluczowe dla rentownosci.', size=10, color=C_GREY, sa=8)

make_table(doc,
    ['Prog pewnosci', 'KNN Mecze', 'KNN Win Rate', 'KNN ROI', 'XGB Mecze', 'XGB Win Rate', 'XGB ROI'],
    [
        ['75%+',  '208', '72.6%', '-655 zl',   '84',  '77.4%', '-418 zl'],
        ['80%+',  '129', '78.3%', '+144 zl',   '48',  '87.5%', '+193 zl'],
        ['82%+',  '76',  '78.9%', '-132 zl',   '35',  '91.4%', '+209 zl'],
        ['85%+',  '76',  '78.9%', '-132 zl',   '21',  '100%',  '+293 zl'],
        ['88%+',  '36',  '77.8%', '-298 zl',   '11',  '100%',  '+117 zl'],
        ['90%+',  '36',  '77.8%', '-298 zl',   '8',   '100%',  '+72 zl'],
        ['93%+',  '11',  '81.8%', '-49 zl',    '6',   '100%',  '+42 zl'],
    ],
    col_widths=[3.2, 2.5, 3.0, 2.8, 2.5, 3.0, 2.8],
    row_colors=[C_LRED, C_LYELL, C_LYELL, C_LGREEN, C_LGREEN, C_LGREEN, C_LGREEN]
)
add_para(doc,
    '\nKNN ma optimum na dokladnie 80% — powyzej 82% ROI spada mimo wyzszego WR (kurs faworyta zbyt niski). '
    'XGBoost: im wyzszy prog tym lepiej — 85%+ = 100% skutecznosc w calym badanym okresie.',
    italic=True, color=C_ORANGE, size=9.5, sa=14)

doc.add_page_break()

# ---- 3. SEGMENTACJA KURSU ----
add_head(doc, '3. Kurs faworyta — strefa zysku vs strefa strat (80%+)', level=2)
add_para(doc, 'Kurs ma ogromny wplyw na wynik — niezaleznie od modelu. Proba: mecze z progiem >= 80%.', size=10, color=C_GREY, sa=8)

make_table(doc,
    ['Kurs faworyta', 'KNN Mecze', 'KNN Win Rate', 'KNN ROI', 'XGB Mecze', 'XGB Win Rate', 'XGB ROI', 'Ocena'],
    [
        ['1.01 - 1.10', '10', '100%',   '+48 zl',   '8',  '100%',  '+51 zl',  'BEZPIECZNA'],
        ['1.10 - 1.20', '18', '94.4%',  '+151 zl',  '15', '93.3%', '+120 zl', 'OPTYMALNA'],
        ['1.20 - 1.30', '35', '74.3%',  '-262 zl',  '14', '71.4%', '-173 zl', 'UNIKAJ'],
        ['1.30 - 1.50', '46', '71.7%',  '-122 zl',  '10', '90.0%', '+195 zl', 'XGB OK, KNN nie'],
        ['1.50+',       '20', '75.0%',  '+329 zl',  '0',  '  --',  '  --',    'Ryzykowna / mala proba'],
    ],
    col_widths=[2.8, 2.4, 3.0, 2.5, 2.4, 3.0, 2.5, 3.5],
    row_colors=[C_LGREEN, C_LGREEN, C_LRED, C_LYELL, C_LYELL]
)
add_para(doc,
    '\nKurs 1.20-1.30 = PULAPKA dla obu modeli. Win Rate spada do 71-74% mimo wysokiej pewnosci. '
    'Model jest w tej strefie systematycznie za pewny siebie.',
    italic=True, color=C_RED, size=9.5, sa=14)

# ---- 4. PRZEGRANE ----
add_head(doc, '4. Dlaczego mecze nie wchodza — analiza przegranych', level=2)
add_head(doc, 'XGBoost — 6 przegranych (prog 80%+, okres 25.03-10.04)', level=3, color=C_RED)

make_table(doc,
    ['Data', 'Mecz', 'Kurs', 'Pewnosc', 'Wynik: typ -> rzeczywisty', 'Glowna przyczyna'],
    [
        ['25.03', 'Lokomotive Leipzig vs Hertha Zehlendorf', '1.25', '82%', '1 -> 2', 'Gosc wygral niespodziewanie'],
        ['04.04', 'Al Hilal vs Al-Taawon',                  '1.22', '82%', '1 -> X', 'REMIS'],
        ['05.04', 'Nasarawa United vs Kun Khalifat',          '1.36', '83%', '1 -> 2', 'Gosc wygral (liga niz.)'],
        ['06.04', 'Catania vs AZ Picerno ASD',               '1.27', '81%', '1 -> 2', 'Strefa 1.20-1.30 (pulapka)'],
        ['07.04', 'Real Espana vs Juticalpa',                '1.18', '80%', '1 -> X', 'REMIS'],
        ['10.04', 'JS Jijel vs JS Bordj Menaiel',            '1.28', '82%', '1 -> 2', 'Strefa 1.20-1.30 (pulapka)'],
    ],
    col_widths=[1.5, 5.8, 1.4, 2.0, 3.8, 4.1],
    hdr_bg=C_RED
)
add_para(doc, '\n5 z 6 przegranych XGB: kurs 1.18-1.36. Przy progu 85%+ XGBoost nie przegral ANI RAZU.', italic=True, color=C_RED, size=9.5, sa=10)

add_head(doc, 'KNN — Analiza 28 przegranych (prog 80%+)', level=3, color=C_RED)
make_table(doc,
    ['Typ przegranej',                                         'Liczba', 'Procent', 'Wniosek'],
    [
        ['REMIS (X) — zamiast wygranej faworyta',             '19', '68%', 'Glowna slabosc modelu KNN'],
        ['Wygral GOSC (2) — zamiast faworyta gospodarza',     '8',  '29%', 'Kursy ok. 1.30-1.65'],
        ['Wygral GOSPODARZ z gola goscia — 1 z faw. goscia', '1',  '3%',  'Pojedynczy przypadek'],
    ],
    col_widths=[8.5, 1.8, 2.2, 5.5],
    row_colors=[C_LRED, C_LYELL, C_LGREY]
)
add_para(doc,
    '\nNajdrozsze przegrane KNN: Liverpool vs Tottenham (1.28, 85%, REMIS), '
    'Panathinaikos vs Panaitolikos (1.30, 90%, REMIS), Admira vs Rapid Wien (1.30, 95%!, gosc wygral).',
    italic=True, color=C_RED, size=9.5, sa=14)

doc.add_page_break()

# ---- 5. TABELA DECYZYJNA ----
add_head(doc, '5. Tabela decyzyjna — CO GRAC, czego unikac', level=2)

make_table(doc,
    ['Scenariusz', 'Prog modelu', 'Kurs faworyta', 'Model', 'Win Rate', 'ROI', 'DECYZJA'],
    [
        ['XGBoost wysoki prog — zloty standard', '>= 85%', '1.05-1.45', 'XGBoost', '100%',   '+293 zl*',  'TYPUJ ZAWSZE'],
        ['Strefa optymalna kursu — oba modele',  '>= 80%', '1.10-1.20', 'KNN/XGB', 'ok. 94%','+150 zl',   'TYPUJ'],
        ['XGBoost z filtrem progu',              '>= 82%', 'dowolny',   'XGBoost', '91.4%',  '+209 zl',   'TYPUJ'],
        ['KNN sweet spot',                       '80-82%', '1.10-1.20', 'KNN',     '94.4%',  '+151 zl',   'TYPUJ'],
        ['Niski prog, bez filtra',               '75-80%', 'dowolny',   'Oba',     '72-77%', 'MINUS',     'NIE TYPUJ'],
        ['Pulapka kursu 1.20-1.30',              'dowolny','1.20-1.30', 'Oba',     '71-74%', 'MINUS',     'UNIKAJ'],
        ['KNN z za wysokim progiem',             '>= 82%', 'dowolny',   'KNN',     '78.9%',  '-132 zl',   'UNIKAJ (paradoks)'],
        ['Szansa remisu > 10%',                  'dowolny','dowolny',   'Oba',     'NISKI',  'MINUS',     'NIE TYPUJ'],
    ],
    col_widths=[5.0, 2.8, 3.0, 2.5, 2.5, 2.3, 3.5],
    hdr_bg=C_NAVYBL,
    row_colors=[C_LGREEN, C_LGREEN, C_LGREEN, C_LGREEN, C_LRED, C_LRED, C_LRED, C_LRED]
)
add_para(doc,
    '\n* XGBoost 85%+ = 21 meczow w badanym okresie (25.03-10.04.2026), 0 przegranych. '
    'Uwaga: mala proba — tylko 16 dni dzialania XGB.',
    italic=True, color=C_GREY, size=9, sa=14)

# ---- 6. STRATEGIE ----
add_head(doc, '6. Trzy strategie — praktyczny dobor', level=2)

make_table(doc,
    ['Strategia', 'Warunek wejscia', 'Opis praktyczny', 'Oczekiwany wynik'],
    [
        ['A — KONSERWATYWNA\n(malo typow, max pewnosc)',
         'XGBoost >= 85%\nKurs: dowolny',
         'Ok. 1 mecz dziennie.\nBrak przegranych w badanym okresie.',
         '100% WR, +293 zl / 21 meczow.\nMoze miec chudsze tygodnie.'],
        ['B — UMIARKOWANA\n(optimum zysku)',
         'KNN lub XGB >= 80%\nKurs: 1.10-1.20',
         'Kilka typow tygodniowo.\nOba modele daja ~94% w tej strefie.',
         'WR ok. 94%\nROI solidnie dodatni.'],
        ['C — Z FILTREM REMISU\n(eliminacja gl. slabosci)',
         'KNN lub XGB >= 80%\n% Remis < 8% w predykcji',
         'Eliminuje 68% przegranych KNN.\nWymaga odczytu kolumny % Remisu [X].',
         'Prognozowana poprawa WR\no 5-10 pp. Brak korekty w kodzie.'],
    ],
    col_widths=[4.0, 3.8, 5.5, 5.5],
    hdr_bg=C_DKGREEN,
    row_colors=[C_LGREEN, C_LGREEN, C_LYELL]
)

doc.add_paragraph()

# Stopka
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
f.paragraph_format.space_before = Pt(20)
rf = f.add_run(
    'Raport wygenerowany automatycznie | '
    'Dane realne: 1 237 meczow KNN (31 dni) + 645 meczow XGBoost (16 dni) | '
    'Praca Inzynierska 2026'
)
rf.font.size = Pt(8); rf.italic = True; rf.font.color.rgb = rgb(C_GREY)

doc.save(OUTPUT_PATH)
print(f'SUKCES — Zapisano: {OUTPUT_PATH}')
